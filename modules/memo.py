"""Deal memo generation — produces Word (.docx) and PDF outputs.

Uses python-docx for Word output (cross-platform, no LibreOffice required).
For PDF, uses ReportLab (pure Python, no system dependencies).

The memo is a single-page summary of the analysis: property, recommendation,
offer terms, action items, and key diagnostics.
"""
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, List


# ---------------------------------------------------------------------------
# Shared formatting helpers
# ---------------------------------------------------------------------------
def fmt_money(x):
    if x is None: return "—"
    try:
        return f"${float(x):,.0f}"
    except (TypeError, ValueError):
        return str(x)

def fmt_pct(x):
    if x is None: return "—"
    try: return f"{float(x):.1%}"
    except (TypeError, ValueError): return str(x)


# ---------------------------------------------------------------------------
# WORD (.docx) generation
# ---------------------------------------------------------------------------
def build_word_memo(prop: Dict, rec: Dict, seller: Dict, rehab_items: List = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # === HEADER ===
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("EXODUS PROPERTY SOLUTIONS")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Acquisitions Deal Memo").font.size = Pt(12)
    sub.runs[0].italic = True

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(datetime.now().strftime("Generated %B %d, %Y")).font.size = Pt(9)

    # === STRATEGY BANNER ===
    doc.add_paragraph()
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = banner.add_run("Recommended Strategy")
    r.bold = True; r.font.size = Pt(10)

    strat_p = doc.add_paragraph()
    strat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = strat_p.add_run(rec.get("strategy", ""))
    r.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    rationale = doc.add_paragraph()
    rationale.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rationale.add_run(rec.get("rationale", ""))
    r.italic = True; r.font.size = Pt(10)

    # === PROPERTY ===
    doc.add_paragraph()
    _add_section(doc, "Property")
    _add_kv_table(doc, [
        ("Address", prop.get("address", "")),
        ("Location", f"{prop.get('city','')}, {prop.get('state','')} {prop.get('zip','')}"),
        ("Beds / Baths / Sqft / Year",
         f"{prop.get('beds','—')} bd / {prop.get('baths','—')} ba / "
         f"{prop.get('sqft',0):,} sf / {prop.get('year','—')}"),
        ("Pool / HOA", f"{prop.get('pool','No')} / {fmt_money(prop.get('hoa',0))}/mo"),
        ("Annual Property Taxes", fmt_money(prop.get("annual_taxes", 0))),
        ("Acquisition Type", prop.get("acquisition_type", "Regular")),
        ("Seller's Asking", fmt_money(prop.get("asking", 0))),
    ])

    # === KEY NUMBERS === (strategy-aware — investor vs MLS vs Pass)
    from modules.strategy import key_numbers_for
    _add_section(doc, "Key Numbers")
    _add_kv_table(doc, key_numbers_for(rec, prop))

    # === REHAB LINE ITEMS ===
    if rehab_items:
        _add_section(doc, "Rehab Line Items")
        rehab_rows = [(label, fmt_money(amount)) for label, amount in rehab_items]
        sub = sum(a for _, a in rehab_items)
        contingency_amt = rec.get("rehab_total", 0) - sub
        rehab_rows.append(("Subtotal", fmt_money(sub)))
        rehab_rows.append((f"Contingency ({'10%' if sub > 50000 else '$5,000 flat'})",
                           fmt_money(contingency_amt)))
        rehab_rows.append(("TOTAL REHAB", fmt_money(rec.get("rehab_total", 0))))
        _add_kv_table(doc, rehab_rows)

    # === GAP ANALYSIS ===
    _add_section(doc, "Asking-MAO Gap")
    _add_kv_table(doc, [
        ("Gap (Asking − Cash MAO)", fmt_money(rec.get("gap", 0))),
        ("Gap Category", rec.get("gap_category", "—")),
        ("Max Asking for Novation", fmt_money(rec.get("novation_max_asking", 0))),
        ("Est. MLS Commission", fmt_money(rec.get("mls_commission_estimate", 0))),
    ])

    # === FINANCIAL PRO-FORMA ===
    # Shown for ALL strategies except MLS Referral (which has its own commission
    # breakdown). For Pass strategies, it shows the math we walked away from so
    # the team can see why and check if a small change would flip the decision.
    strat = rec.get("strategy", "")
    show_proforma = "MLS" not in strat
    is_pass = "Pass" in strat or strat.startswith("NO-GO")
    if show_proforma:
        section_title = ("Financial Pro-Forma — Why We Passed"
                         if is_pass else "Financial Pro-Forma")
        _add_section(doc, section_title)
        # For Pass strategies, default to Rehab-style math to show what the deal
        # would look like if we'd taken it down.
        if "Rehab" in strat or is_pass:
            # Rehab pro-forma: full acquisition, rehab, hold, retail sale
            purchase_price = rec.get("likely_purchase_price", rec.get("cash_offer", 0))
            purchase_label = ("Purchase Price (at Asking)"
                              if rec.get("net_profit_at_asking") is not None
                              else "Purchase Price (Cash MAO)")
            monthly_holding_total = rec.get('monthly_holding', 0)
            ins_line = rec.get('monthly_insurance', 0)
            tax_line = rec.get('monthly_taxes', 0)
            arv = rec.get("arv", 0)
            bc_closing = rec.get("sale_closing_costs", 0)
            ab_closing = rec.get("purchase_closing_costs", 0)
            rehab = rec.get("rehab_total", 0)
            holding = rec.get("total_holding", 0)
            com = rec.get("cost_of_money", 0)
            # v24 itemized closing breakdowns (fall back to single line if absent)
            ab_items = (rec.get("ab_closing_itemized") or {}).get("line_items", [])
            bc_items = (rec.get("bc_closing_itemized") or {}).get("line_items", [])
            bc_comm_note = (rec.get("bc_closing_itemized") or {}).get(
                "commission_note", "")
            household_adj_net = rec.get("net_profit_household_adjusted")
            listing_recoverable = rec.get("commission_listing_recoverable", 0)
            # Net Acquisition+Project Cost = TPC minus BC (BC is shown on sale side)
            net_acq_cost = purchase_price + ab_closing + rehab + holding + com
            net_sale_proceeds = arv - bc_closing

            proforma_rows = [
                ("SALE SIDE (B→C)", ""),
                ("  Expected Sale Price (ARV)", fmt_money(arv)),
                ("  Less: BC Closing Costs — itemized", ""),
            ]
            # Itemized BC breakdown
            if bc_items:
                for label, amt in bc_items:
                    proforma_rows.append(
                        (f"      {label}", f"({fmt_money(amt)})"))
                if bc_comm_note:
                    proforma_rows.append(("      " + bc_comm_note, ""))
            else:
                proforma_rows.append(
                    (f"      BC Closing ({rec.get('sale_closing_pct', 0.07):.1%})",
                     f"({fmt_money(bc_closing)})"))
            proforma_rows.extend([
                ("  Total BC Closing Costs", f"({fmt_money(bc_closing)})"),
                ("  Net Sale Proceeds", fmt_money(net_sale_proceeds)),
                ("", ""),
                ("ACQUISITION & PROJECT (A→B + Hold)", ""),
                (f"  {purchase_label}", fmt_money(purchase_price)),
                ("  Plus: AB Closing Costs — itemized", ""),
            ])
            # Itemized AB breakdown
            if ab_items:
                for label, amt in ab_items:
                    proforma_rows.append((f"      {label}", fmt_money(amt)))
            else:
                proforma_rows.append(
                    (f"      AB Closing ({rec.get('purchase_closing_pct', 0.04):.1%})",
                     fmt_money(ab_closing)))
            proforma_rows.extend([
                ("  Total AB Closing Costs", fmt_money(ab_closing)),
                ("  Plus: Total Rehab", fmt_money(rehab)),
                (f"  Plus: Holding Costs ({rec.get('loan_duration_months', 6)} mo × "
                 f"{fmt_money(monthly_holding_total)}/mo)",
                 fmt_money(holding)),
                (f"      — incl. Insurance {fmt_money(ins_line)}/mo + "
                 f"Taxes {fmt_money(tax_line)}/mo", ""),
                (f"  Plus: Cost of Money ({rec.get('ltc', 0.9):.0%} LTC, "
                 f"loan {fmt_money(rec.get('likely_loan', 0))} @ "
                 f"{rec.get('interest_rate', 0.10):.1%})",
                 fmt_money(com)),
                ("  Total Acquisition + Project Cost", fmt_money(net_acq_cost)),
                ("", ""),
                ("BOTTOM LINE", ""),
                ("  Net Sale Proceeds − Total Cost = Net Profit (LLC view)",
                 fmt_money(rec.get("net_profit", 0))),
                ("  Projected ROI", fmt_pct(rec.get("roi", 0))),
            ])
            # Household-adjusted view — only meaningful if listing commission > 0
            if listing_recoverable and listing_recoverable > 0 and household_adj_net is not None:
                proforma_rows.extend([
                    ("", ""),
                    ("  * Household-Adjusted Net (adds listing agent commission back)",
                     fmt_money(household_adj_net)),
                    (f"    Listing commission of {fmt_money(listing_recoverable)} "
                     "is paid to a licensed family member and effectively "
                     "returns to the household. Still counted as a deal cost "
                     "for conservative underwriting.", ""),
                ])
        elif rec.get("proforma_kind") == "assignment":
            # Wholesale Assignment — we never close, profit = the fee
            proforma_rows = [
                ("STRATEGY", "Assignment — no closings, no rehab on our side"),
                ("", ""),
                ("ASSIGNMENT FEE", ""),
                (f"  Spread we collect from end buyer",
                 fmt_money(rec.get("target_assignment_fee") or rec.get("net_profit", 0))),
                ("", ""),
                ("OUR COSTS", ""),
                ("  AB Closing (we don't close)", "$0"),
                ("  BC Closing (end buyer pays)", "$0"),
                ("  Holding / COM", "$0"),
                ("", ""),
                ("BOTTOM LINE", ""),
                ("  Net Profit", fmt_money(rec.get("net_profit", 0))),
            ]
        elif rec.get("proforma_kind") == "novation":
            # Novation — list at ARV; profit = ARV × (1 - retail%) − seller's net − rehab − holding
            proforma_rows = [
                ("STRATEGY", "Novation — control via agreement, sell at retail"),
                ("", ""),
                ("SALE SIDE (Retail listing)", ""),
                ("  Expected Sale Price (ARV)", fmt_money(rec.get("arv", 0))),
                (f"  Less: Retail Costs ({rec.get('sale_closing_pct', 0.09)*100:.1f}% of ARV)",
                 f"({fmt_money(rec.get('sale_closing_costs', 0))})"),
                ("", ""),
                ("LESS: SELLER'S REQUIRED NET", ""),
                ("  Benchmark (asking or required net)",
                 f"({fmt_money(rec.get('benchmark', 0))})"),
                ("", ""),
                ("LESS: OUR COSTS", ""),
                ("  Rehab (if any)", f"({fmt_money(rec.get('rehab_total', 0))})"),
                ("  Novation Holding Costs",
                 f"({fmt_money(rec.get('total_holding', 0))})"),
                ("", ""),
                ("BOTTOM LINE", ""),
                ("  Net Profit", fmt_money(rec.get("net_profit", 0))),
            ]
        else:
            # Double Close pro-forma — buy from seller, immediate resell to end buyer
            # at THEIR Cash MAO. We do NOT do the rehab.
            our_buy_price = rec.get("likely_purchase_price", rec.get("cash_offer", 0))
            end_buyer_price = rec.get("cash_offer", 0)
            ab_closing_dc = rec.get("purchase_closing_costs", 0)
            bc_closing_dc = rec.get("sale_closing_costs", 0)
            transactional_fee = rec.get("cost_of_money", 0)
            ab_pct = rec.get("purchase_closing_pct", 0.04)
            bc_pct = rec.get("sale_closing_pct", 0.02)
            total_dc_costs = ab_closing_dc + bc_closing_dc + transactional_fee
            proforma_rows = [
                ("STRATEGY", "Double Close — buy and immediately resell; end buyer does the rehab"),
                ("", ""),
                ("SELL SIDE (B→C, same day)", ""),
                ("  End-Buyer Price (their Cash MAO)", fmt_money(end_buyer_price)),
                (f"  Less: BC Closing Costs ({bc_pct*100:.1f}%)",
                 f"({fmt_money(bc_closing_dc)})"),
                ("  Net from End Buyer", fmt_money(end_buyer_price - bc_closing_dc)),
                ("", ""),
                ("BUY SIDE (A→B)", ""),
                ("  Our Contract Price (to Seller)", fmt_money(our_buy_price)),
                (f"  Plus: AB Closing Costs ({ab_pct*100:.1f}%)",
                 fmt_money(ab_closing_dc)),
                ("  Plus: Transactional Funding (1%, same-day bridge)",
                 fmt_money(transactional_fee)),
                ("  Total Our Costs (out of pocket)",
                 fmt_money(our_buy_price + ab_closing_dc + transactional_fee)),
                ("", ""),
                ("BOTTOM LINE", ""),
                ("  Gross Spread (B − A)",
                 fmt_money(end_buyer_price - our_buy_price)),
                ("  Less: All Closing & Funding Costs",
                 f"({fmt_money(total_dc_costs)})"),
                ("  Net Profit", fmt_money(rec.get("net_profit", 0))),
            ]
        _add_kv_table(doc, proforma_rows)

    # === OFFER TERMS (skip for MLS/Pass) ===
    if not (strat.startswith("Pass") or strat == "NO-GO — Pass" or strat == "MLS Referral"):
        _add_section(doc, "Offer Terms")
        rows = [
            ("Opening Offer", fmt_money(rec.get("opening_offer", 0))),
            ("Walk-Away (MAO)", fmt_money(rec.get("walk_away", 0))),
            ("Stretch Ceiling", fmt_money(rec.get("stretch_ceiling", 0))),
        ]
        if rec.get("target_assignment_fee") is not None:
            rows.append(("Target Assignment Fee", fmt_money(rec["target_assignment_fee"])))
            if rec.get("fat_fee_note"):
                rows.append(("Fee Notes", rec["fat_fee_note"]))
        ct = rec.get("contract_terms", {})
        rows.extend([
            ("Offer Type", ct.get("offer_type", "")),
            ("Earnest Money", fmt_money(ct.get("earnest_money", 0))),
            ("Inspection Period", ct.get("inspection_period", "")),
            ("Close Date", ct.get("close_date", "")),
            ("Assignment Language", ct.get("assignment_language", "")),
        ])
        _add_kv_table(doc, rows)

    # === SELLER INFO ===
    _add_section(doc, "Seller & Loan Info")
    _add_kv_table(doc, [
        ("1st Mortgage Balance", fmt_money(seller.get("mtg1", 0))),
        ("2nd / HELOC Balance", fmt_money(seller.get("mtg2", 0))),
        ("Other Liens", fmt_money(seller.get("other_liens", 0))),
        ("Equity Position", fmt_money(rec.get("equity", 0))),
        ("Payment Status", seller.get("payment_status", "—")),
        ("Seller's Required Net", fmt_money(seller.get("required_net", 0))),
        ("Timeline (days)", str(seller.get("timeline", "—"))),
        ("Reason for Selling", seller.get("reason", "—")),
        ("Occupancy", seller.get("occupancy", "—")),
        ("Condition Confirmed", seller.get("condition_confirmed", "—")),
        ("Open to MLS Listing", seller.get("open_to_mls", "—")),
    ])

    # === DISPOSITION ===
    _add_section(doc, "Disposition Plan")
    doc.add_paragraph(rec.get("disposition", ""))

    # === ACTION ITEMS ===
    actions = rec.get("action_items", [])
    if actions:
        _add_section(doc, "Action Items")
        for i, action in enumerate(actions, 1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(action)

    # === FOOTER ===
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.add_run("Generated by the Exodus Underwriting Tool").font.size = Pt(8)
    foot.runs[0].italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_section(doc, title: str):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)


def _add_kv_table(doc, rows: List):
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (label, val) in enumerate(rows):
        c1 = table.rows[i].cells[0]
        c1.width = Inches(2.5)
        c1.text = str(label)
        c1.paragraphs[0].runs[0].font.size = Pt(10)
        c1.paragraphs[0].runs[0].bold = True
        # Light gray fill
        tcPr = c1._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F2F2F2')
        tcPr.append(shd)

        c2 = table.rows[i].cells[1]
        c2.width = Inches(4.5)
        c2.text = str(val)
        c2.paragraphs[0].runs[0].font.size = Pt(10)


# ---------------------------------------------------------------------------
# PDF generation (ReportLab — pure Python, no system deps)
# ---------------------------------------------------------------------------
def build_pdf_memo(prop: Dict, rec: Dict, seller: Dict, rehab_items: List = None) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, KeepTogether)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                            leftMargin=0.6*inch, rightMargin=0.6*inch,
                            topMargin=0.5*inch, bottomMargin=0.5*inch)

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle("Title", parent=styles["Heading1"],
                             fontName="Helvetica-Bold", fontSize=16,
                             textColor=colors.HexColor("#1F4E78"),
                             alignment=TA_CENTER, spaceAfter=4)
    sub_s = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=11,
                           textColor=colors.HexColor("#666666"),
                           alignment=TA_CENTER, spaceAfter=14)
    section_s = ParagraphStyle("Section", parent=styles["Heading2"],
                               fontName="Helvetica-Bold", fontSize=12,
                               textColor=colors.HexColor("#1F4E78"),
                               spaceBefore=12, spaceAfter=6)
    strat_s = ParagraphStyle("Strat", parent=styles["Normal"], fontSize=18,
                             fontName="Helvetica-Bold",
                             textColor=colors.HexColor("#1F4E78"),
                             alignment=TA_CENTER, spaceAfter=4)
    rat_s = ParagraphStyle("Rat", parent=styles["Normal"], fontSize=10,
                           textColor=colors.HexColor("#333333"),
                           alignment=TA_CENTER, spaceAfter=10, italic=True)
    normal_s = ParagraphStyle("Norm", parent=styles["Normal"], fontSize=10)

    story = []
    story.append(Paragraph("EXODUS PROPERTY SOLUTIONS", title_s))
    story.append(Paragraph(
        f"Acquisitions Deal Memo &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"{datetime.now().strftime('%B %d, %Y')}", sub_s))

    # Strategy banner
    story.append(Paragraph("Recommended Strategy", ParagraphStyle(
        "ResL", parent=normal_s, alignment=TA_CENTER, fontSize=9,
        textColor=colors.HexColor("#666666"))))
    story.append(Paragraph(rec.get("strategy", ""), strat_s))
    story.append(Paragraph(rec.get("rationale", ""), rat_s))

    def kv_table(rows):
        t = Table([[Paragraph(str(k), normal_s), Paragraph(str(v), normal_s)] for k, v in rows],
                  colWidths=[2.4*inch, 4.7*inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (0,-1), colors.HexColor("#F2F2F2")),
            ("FONTNAME", (0,0), (0,-1), "Helvetica-Bold"),
            ("INNERGRID", (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
            ("BOX", (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 6),
            ("RIGHTPADDING", (0,0), (-1,-1), 6),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ]))
        return t

    # Property
    story.append(Paragraph("Property", section_s))
    story.append(kv_table([
        ("Address", prop.get("address", "")),
        ("Location", f"{prop.get('city','')}, {prop.get('state','')} {prop.get('zip','')}"),
        ("Beds / Baths / Sqft / Year",
         f"{prop.get('beds','—')} bd / {prop.get('baths','—')} ba / "
         f"{prop.get('sqft',0):,} sf / {prop.get('year','—')}"),
        ("Pool / HOA", f"{prop.get('pool','No')} / {fmt_money(prop.get('hoa',0))}/mo"),
        ("Annual Property Taxes", fmt_money(prop.get("annual_taxes", 0))),
        ("Acquisition Type", prop.get("acquisition_type", "Regular")),
        ("Seller's Asking", fmt_money(prop.get("asking", 0))),
    ]))

    # Key Numbers — strategy-aware (investor vs MLS vs Pass)
    from modules.strategy import key_numbers_for
    story.append(Paragraph("Key Numbers", section_s))
    story.append(kv_table(key_numbers_for(rec, prop)))

    # Rehab Line Items
    if rehab_items:
        story.append(Paragraph("Rehab Line Items", section_s))
        rehab_rows = [(label, fmt_money(amount)) for label, amount in rehab_items]
        sub = sum(a for _, a in rehab_items)
        contingency_amt = rec.get("rehab_total", 0) - sub
        rehab_rows.append(("Subtotal", fmt_money(sub)))
        rehab_rows.append((f"Contingency ({'10%' if sub > 50000 else '$5,000 flat'})",
                           fmt_money(contingency_amt)))
        rehab_rows.append(("<b>TOTAL REHAB</b>", f"<b>{fmt_money(rec.get('rehab_total', 0))}</b>"))
        story.append(kv_table(rehab_rows))

    # Gap
    story.append(Paragraph("Asking-MAO Gap", section_s))
    story.append(kv_table([
        ("Gap (Asking − Cash MAO)", fmt_money(rec.get("gap", 0))),
        ("Gap Category", rec.get("gap_category", "—")),
        ("Max Asking for Novation", fmt_money(rec.get("novation_max_asking", 0))),
        ("Est. MLS Commission", fmt_money(rec.get("mls_commission_estimate", 0))),
    ]))

    # Financial Pro-Forma — shown for everything except MLS Referral
    strat = rec.get("strategy", "")
    show_proforma = "MLS" not in strat
    is_pass = "Pass" in strat or strat.startswith("NO-GO")
    if show_proforma:
        section_title = ("Financial Pro-Forma — Why We Passed"
                         if is_pass else "Financial Pro-Forma")
        story.append(Paragraph(section_title, section_s))
        if "Rehab" in strat or is_pass:
            purchase_price = rec.get("likely_purchase_price", rec.get("cash_offer", 0))
            purchase_label = ("Purchase Price (at Asking)"
                              if rec.get("net_profit_at_asking") is not None
                              else "Purchase Price (Cash MAO)")
            ins_line = rec.get('monthly_insurance', 0)
            tax_line = rec.get('monthly_taxes', 0)
            arv = rec.get("arv", 0)
            bc_closing = rec.get("sale_closing_costs", 0)
            ab_closing = rec.get("purchase_closing_costs", 0)
            rehab = rec.get("rehab_total", 0)
            holding = rec.get("total_holding", 0)
            com = rec.get("cost_of_money", 0)
            # v24 itemized closing breakdowns
            ab_items = (rec.get("ab_closing_itemized") or {}).get("line_items", [])
            bc_items = (rec.get("bc_closing_itemized") or {}).get("line_items", [])
            bc_comm_note = (rec.get("bc_closing_itemized") or {}).get(
                "commission_note", "")
            household_adj_net = rec.get("net_profit_household_adjusted")
            listing_recoverable = rec.get("commission_listing_recoverable", 0)
            # Net Acquisition+Project Cost excludes BC (BC is on sale side)
            net_acq_cost = purchase_price + ab_closing + rehab + holding + com
            net_sale_proceeds = arv - bc_closing
            proforma_rows = [
                ("<b>SALE SIDE (B→C)</b>", ""),
                ("&nbsp;&nbsp;Expected Sale Price (ARV)", fmt_money(arv)),
                ("&nbsp;&nbsp;<i>Less: BC Closing Costs — itemized</i>", ""),
            ]
            if bc_items:
                for label, amt in bc_items:
                    proforma_rows.append(
                        (f"&nbsp;&nbsp;&nbsp;&nbsp;{label}",
                         f"({fmt_money(amt)})"))
                if bc_comm_note:
                    proforma_rows.append(
                        (f"&nbsp;&nbsp;&nbsp;&nbsp;<i>{bc_comm_note}</i>", ""))
            else:
                proforma_rows.append(
                    (f"&nbsp;&nbsp;&nbsp;&nbsp;BC Closing "
                     f"({rec.get('sale_closing_pct', 0.07):.1%})",
                     f"({fmt_money(bc_closing)})"))
            proforma_rows.extend([
                ("&nbsp;&nbsp;<b>Total BC Closing Costs</b>",
                 f"<b>({fmt_money(bc_closing)})</b>"),
                ("&nbsp;&nbsp;Net Sale Proceeds", fmt_money(net_sale_proceeds)),
                ("<b>ACQUISITION &amp; PROJECT (A→B + Hold)</b>", ""),
                (f"&nbsp;&nbsp;{purchase_label}", fmt_money(purchase_price)),
                ("&nbsp;&nbsp;<i>Plus: AB Closing Costs — itemized</i>", ""),
            ])
            if ab_items:
                for label, amt in ab_items:
                    proforma_rows.append(
                        (f"&nbsp;&nbsp;&nbsp;&nbsp;{label}", fmt_money(amt)))
            else:
                proforma_rows.append(
                    (f"&nbsp;&nbsp;&nbsp;&nbsp;AB Closing "
                     f"({rec.get('purchase_closing_pct', 0.04):.1%})",
                     fmt_money(ab_closing)))
            proforma_rows.extend([
                ("&nbsp;&nbsp;<b>Total AB Closing Costs</b>",
                 f"<b>{fmt_money(ab_closing)}</b>"),
                ("&nbsp;&nbsp;Plus: Total Rehab", fmt_money(rehab)),
                (f"&nbsp;&nbsp;Plus: Holding Costs ({rec.get('loan_duration_months', 6)} mo × "
                 f"{fmt_money(rec.get('monthly_holding', 0))}/mo)",
                 fmt_money(holding)),
                (f"&nbsp;&nbsp;&nbsp;&nbsp;<i>— incl. Insurance {fmt_money(ins_line)}/mo "
                 f"+ Taxes {fmt_money(tax_line)}/mo</i>", ""),
                (f"&nbsp;&nbsp;Plus: Cost of Money ({rec.get('ltc', 0.9):.0%} LTC, "
                 f"loan {fmt_money(rec.get('likely_loan', 0))} @ "
                 f"{rec.get('interest_rate', 0.10):.1%})",
                 fmt_money(com)),
                ("&nbsp;&nbsp;Total Acquisition + Project Cost",
                 fmt_money(net_acq_cost)),
                ("<b>BOTTOM LINE</b>", ""),
                ("&nbsp;&nbsp;Net Sale Proceeds − Total Cost", ""),
                ("&nbsp;&nbsp;<b>Net Profit (LLC view)</b>",
                 f"<b>{fmt_money(rec.get('net_profit', 0))}</b>"),
                ("&nbsp;&nbsp;Projected ROI", fmt_pct(rec.get("roi", 0))),
            ])
            # Household-adjusted view
            if listing_recoverable and listing_recoverable > 0 and household_adj_net is not None:
                proforma_rows.extend([
                    ("&nbsp;&nbsp;* <b>Household-Adjusted Net</b> (listing commission added back)",
                     f"<b>{fmt_money(household_adj_net)}</b>"),
                    (f"&nbsp;&nbsp;&nbsp;&nbsp;<i>Listing commission "
                     f"{fmt_money(listing_recoverable)} is paid to a licensed "
                     "family member and returns to the household. Still counted "
                     "as a deal cost for conservative LLC underwriting.</i>", ""),
                ])
        elif rec.get("proforma_kind") == "assignment":
            proforma_rows = [
                ("<b>STRATEGY</b>", "Assignment — no closings, no rehab on our side"),
                ("<b>ASSIGNMENT FEE</b>", ""),
                ("&nbsp;&nbsp;Spread we collect from end buyer",
                 fmt_money(rec.get("target_assignment_fee") or rec.get("net_profit", 0))),
                ("<b>OUR COSTS</b>", ""),
                ("&nbsp;&nbsp;AB Closing (we don't close)", "$0"),
                ("&nbsp;&nbsp;BC Closing (end buyer pays)", "$0"),
                ("&nbsp;&nbsp;Holding / COM", "$0"),
                ("<b>BOTTOM LINE</b>", ""),
                ("&nbsp;&nbsp;<b>Net Profit</b>",
                 f"<b>{fmt_money(rec.get('net_profit', 0))}</b>"),
            ]
        elif rec.get("proforma_kind") == "novation":
            proforma_rows = [
                ("<b>STRATEGY</b>", "Novation — control via agreement, sell at retail"),
                ("<b>SALE SIDE (Retail listing)</b>", ""),
                ("&nbsp;&nbsp;Expected Sale Price (ARV)", fmt_money(rec.get("arv", 0))),
                (f"&nbsp;&nbsp;Less: Retail Costs ({rec.get('sale_closing_pct', 0.09)*100:.1f}% of ARV)",
                 f"({fmt_money(rec.get('sale_closing_costs', 0))})"),
                ("<b>LESS: SELLER'S REQUIRED NET</b>", ""),
                ("&nbsp;&nbsp;Benchmark (asking or required net)",
                 f"({fmt_money(rec.get('benchmark', 0))})"),
                ("<b>LESS: OUR COSTS</b>", ""),
                ("&nbsp;&nbsp;Rehab (if any)",
                 f"({fmt_money(rec.get('rehab_total', 0))})"),
                ("&nbsp;&nbsp;Novation Holding Costs",
                 f"({fmt_money(rec.get('total_holding', 0))})"),
                ("<b>BOTTOM LINE</b>", ""),
                ("&nbsp;&nbsp;<b>Net Profit</b>",
                 f"<b>{fmt_money(rec.get('net_profit', 0))}</b>"),
            ]
        else:
            # Double Close
            our_buy_price = rec.get("likely_purchase_price", rec.get("cash_offer", 0))
            end_buyer_price = rec.get("cash_offer", 0)
            ab_closing_dc = rec.get("purchase_closing_costs", 0)
            bc_closing_dc = rec.get("sale_closing_costs", 0)
            transactional_fee = rec.get("cost_of_money", 0)
            ab_pct = rec.get("purchase_closing_pct", 0.04)
            bc_pct = rec.get("sale_closing_pct", 0.02)
            total_dc_costs = ab_closing_dc + bc_closing_dc + transactional_fee
            proforma_rows = [
                ("<b>STRATEGY</b>", "Double Close — buy and immediately resell; end buyer does the rehab"),
                ("<b>SELL SIDE (B→C, same day)</b>", ""),
                ("&nbsp;&nbsp;End-Buyer Price (their Cash MAO)",
                 fmt_money(end_buyer_price)),
                (f"&nbsp;&nbsp;Less: BC Closing Costs ({bc_pct*100:.1f}%)",
                 f"({fmt_money(bc_closing_dc)})"),
                ("&nbsp;&nbsp;Net from End Buyer",
                 fmt_money(end_buyer_price - bc_closing_dc)),
                ("<b>BUY SIDE (A→B)</b>", ""),
                ("&nbsp;&nbsp;Our Contract Price (to Seller)",
                 fmt_money(our_buy_price)),
                (f"&nbsp;&nbsp;Plus: AB Closing Costs ({ab_pct*100:.1f}%)",
                 fmt_money(ab_closing_dc)),
                ("&nbsp;&nbsp;Plus: Transactional Funding (1%, same-day bridge)",
                 fmt_money(transactional_fee)),
                ("&nbsp;&nbsp;Total Our Costs",
                 fmt_money(our_buy_price + ab_closing_dc + transactional_fee)),
                ("<b>BOTTOM LINE</b>", ""),
                ("&nbsp;&nbsp;Gross Spread (B − A)",
                 fmt_money(end_buyer_price - our_buy_price)),
                ("&nbsp;&nbsp;Less: All Closing &amp; Funding Costs",
                 f"({fmt_money(total_dc_costs)})"),
                ("&nbsp;&nbsp;<b>Net Profit</b>",
                 f"<b>{fmt_money(rec.get('net_profit', 0))}</b>"),
            ]
        story.append(kv_table(proforma_rows))

    # Offer Terms (conditional)
    if not (strat.startswith("Pass") or strat == "NO-GO — Pass" or strat == "MLS Referral"):
        story.append(Paragraph("Offer Terms", section_s))
        offer_rows = [
            ("Opening Offer", fmt_money(rec.get("opening_offer", 0))),
            ("Walk-Away (MAO)", fmt_money(rec.get("walk_away", 0))),
            ("Stretch Ceiling", fmt_money(rec.get("stretch_ceiling", 0))),
        ]
        if rec.get("target_assignment_fee") is not None:
            offer_rows.append(("Target Assignment Fee",
                               fmt_money(rec["target_assignment_fee"])))
        ct = rec.get("contract_terms", {})
        offer_rows.extend([
            ("Offer Type", ct.get("offer_type", "")),
            ("Earnest Money", fmt_money(ct.get("earnest_money", 0))),
            ("Inspection Period", ct.get("inspection_period", "")),
            ("Close Date", ct.get("close_date", "")),
            ("Assignment Language", ct.get("assignment_language", "")),
        ])
        story.append(kv_table(offer_rows))

    # Seller Info
    story.append(Paragraph("Seller &amp; Loan Info", section_s))
    story.append(kv_table([
        ("1st Mortgage Balance", fmt_money(seller.get("mtg1", 0))),
        ("2nd / HELOC Balance", fmt_money(seller.get("mtg2", 0))),
        ("Other Liens", fmt_money(seller.get("other_liens", 0))),
        ("Equity Position", fmt_money(rec.get("equity", 0))),
        ("Payment Status", seller.get("payment_status", "—")),
        ("Seller's Required Net", fmt_money(seller.get("required_net", 0))),
        ("Timeline (days)", str(seller.get("timeline", "—"))),
        ("Reason for Selling", seller.get("reason", "—")),
        ("Occupancy", seller.get("occupancy", "—")),
        ("Condition Confirmed", seller.get("condition_confirmed", "—")),
        ("Open to MLS Listing", seller.get("open_to_mls", "—")),
    ]))

    # Disposition
    story.append(Paragraph("Disposition Plan", section_s))
    story.append(Paragraph(rec.get("disposition", ""), normal_s))

    # Action Items
    actions = rec.get("action_items", [])
    if actions:
        story.append(Paragraph("Action Items", section_s))
        for i, a in enumerate(actions, 1):
            story.append(Paragraph(f"{i}. {a}", normal_s))

    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "<i>Generated by the Exodus Underwriting Tool</i>",
        ParagraphStyle("Foot", parent=normal_s, fontSize=8,
                       textColor=colors.HexColor("#999999"), alignment=TA_CENTER)))

    doc.build(story)
    return buf.getvalue()
